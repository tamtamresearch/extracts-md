#!/usr/bin/env python3
"""Convert ISO/CEN/EN 'Extract' .docx files to Markdown for TC204 publishing.

Features:
  * YAML front-matter (standard, name/name_1..n, published/edition/pages, annotation, note)
  * Body from "Introduction" onward, H1 demoted to H2 (Heading n -> level n+1)
  * Embedded figures extracted to <docname>/fig-N.<ext>; EMF/WMF converted to PNG
  * Word tables -> HTML <table> with colspan/rowspan (merged cells preserved)
  * Czech custom styles mapped (Text normy, Seznam v normě, Poznámka, NadpisTabObr, ...)
"""
import glob
import html as _html
import os
import re
import shutil
import subprocess
import sys

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

P_TAG = qn("w:p")
TBL_TAG = qn("w:tbl")

DASH_SPLIT = re.compile(r"\s*[–—]\s*")
CAPTION_RE = re.compile(
    r"Published\s+(?:in\s+)?(?P<year>\d{4})"
    r"(?:\s*\(Edition\s+(?P<edition>\d+)\))?"
    r".*?(?P<pages>\d+)\s+pages",
    re.IGNORECASE,
)
CAPTION_STYLES = {"Caption", "NadpisTabObr"}
LIST_STYLES = {"List Paragraph", "Seznam v normě"}
NOTE_STYLES = {"Poznámka"}
CT_EXT = {
    "image/png": "png", "image/jpeg": "jpg", "image/gif": "gif",
    "image/tiff": "tif", "image/bmp": "bmp",
    "image/x-emf": "emf", "image/emf": "emf",
    "image/x-wmf": "wmf", "image/wmf": "wmf",
}
VECTOR = {"emf", "wmf"}


def _trim_whitespace(png, pad=12):
    """Crop white margins from a rendered vector image."""
    try:
        from PIL import Image, ImageChops
    except ImportError:
        return
    im = Image.open(png).convert("RGB")
    bg = Image.new("RGB", im.size, (255, 255, 255))
    bbox = ImageChops.difference(im, bg).getbbox()
    if not bbox:
        return
    l, t, r, b = bbox
    l, t = max(0, l - pad), max(0, t - pad)
    r, b = min(im.width, r + pad), min(im.height, b + pad)
    im.crop((l, t, r, b)).save(png)


def clean(text: str) -> str:
    text = text.replace("\xa0", " ").replace(" ", " ").replace(" ", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def yaml_q(s: str) -> str:
    s = re.sub(r"\s+", " ", s.replace("\n", " ")).strip()
    return '"' + s.replace("\\", "\\\\").replace('"', '\\"') + '"'


# ---------- metadata ----------

def parse_name_table(table):
    vals = []
    for c in table.rows[0].cells:
        t = clean(c.text)
        if t and t not in vals:
            vals.append(t)
    vals.sort(key=len)
    standard = vals[0]
    raw = re.sub(r"\s*\n\s*", " ", vals[-1])
    parts = [p.strip() for p in DASH_SPLIT.split(raw) if p.strip()]
    name = re.sub(r"\s+", " ", " – ".join(parts))
    return standard, name, parts


# ---------- images ----------

def find_soffice():
    """Locate the LibreOffice binary across platforms.

    macOS Homebrew cask provides `soffice`; Linux packages provide
    `libreoffice` (and usually `soffice` too). Override with $SOFFICE_BIN.
    """
    candidates = [os.environ.get("SOFFICE_BIN"), "soffice", "libreoffice"]
    for c in candidates:
        if c and shutil.which(c):
            return shutil.which(c)
    for p in ("/Applications/LibreOffice.app/Contents/MacOS/soffice",
              "/opt/homebrew/bin/soffice", "/usr/local/bin/soffice"):
        if os.path.exists(p):
            return p
    raise RuntimeError(
        "LibreOffice not found (looked for 'soffice'/'libreoffice'). "
        "Install it (macOS: `brew install --cask libreoffice`) or set "
        "$SOFFICE_BIN to the binary path. Needed to rasterise EMF/WMF figures."
    )


_SOFFICE = None


class ImgState:
    def __init__(self, doc, assetdir, docname):
        self.doc = doc
        self.assetdir = assetdir
        self.docname = docname
        self.cache = {}     # rid -> relative path
        self.n = 0

    def save(self, rid):
        if rid in self.cache:
            return self.cache[rid]
        part = self.doc.part.related_parts[rid]
        ext = CT_EXT.get(part.content_type, "png")
        self.n += 1
        base = f"fig-{self.n}"
        os.makedirs(self.assetdir, exist_ok=True)
        # Stage in /tmp (mounts disallow file deletion); copy final file to mount.
        tmp = os.path.join("/tmp", "extract_assets")
        os.makedirs(tmp, exist_ok=True)
        staged = os.path.join(tmp, base + "." + ext)
        with open(staged, "wb") as fh:
            fh.write(part.blob)
        if ext in VECTOR:
            global _SOFFICE
            if _SOFFICE is None:
                _SOFFICE = find_soffice()
            profile = os.path.join(tmp, "lo_profile")
            subprocess.run(
                [_SOFFICE, "--headless",
                 "-env:UserInstallation=file://" + profile,
                 "--convert-to", "png", "--outdir", tmp, staged],
                check=False,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
            png = os.path.join(tmp, base + ".png")
            if os.path.exists(png):
                _trim_whitespace(png)
                staged, ext = png, "png"
        final = os.path.join(self.assetdir, base + "." + ext)
        shutil.copyfile(staged, final)
        rel = f"{base}.{ext}"          # md lives in same folder as figures
        self.cache[rid] = rel
        return rel


def para_blip_rids(el):
    return [b.get(qn("r:embed")) for b in el.findall(".//" + qn("a:blip"))
            if b.get(qn("r:embed"))]


# ---------- tables ----------

def cell_html(tc, imgs: ImgState) -> str:
    chunks = []
    for p in tc.findall(qn("w:p")):
        for rid in para_blip_rids(p):
            chunks.append(f'<img src="{imgs.save(rid)}" alt="">')
        txt = clean("".join(t.text or "" for t in p.iter(qn("w:t"))))
        if txt:
            chunks.append(_html.escape(txt))
    return "<br>".join(chunks)


def table_html(table, imgs: ImgState) -> str:
    trs = table._tbl.findall(qn("w:tr"))
    parsed = []
    for tr in trs:
        col, cells = 0, []
        for tc in tr.findall(qn("w:tc")):
            gs = tc.find(".//" + qn("w:gridSpan"))
            span = int(gs.get(qn("w:val"))) if gs is not None else 1
            vm = tc.find(".//" + qn("w:vMerge"))
            v = (vm.get(qn("w:val")) or "continue") if vm is not None else None
            cells.append({"tc": tc, "col": col, "span": span, "vm": v})
            col += span
        parsed.append(cells)

    nrows = len(parsed)
    for ri, row in enumerate(parsed):
        for c in row:
            if c["vm"] == "continue":
                c["render"] = False
                continue
            c["render"] = True
            rs = 1
            if c["vm"] == "restart":
                for rj in range(ri + 1, nrows):
                    if any(x["col"] == c["col"] and x["vm"] == "continue"
                           for x in parsed[rj]):
                        rs += 1
                    else:
                        break
            c["rowspan"] = rs

    out = ["<table>"]
    for ri, row in enumerate(parsed):
        out.append("  <tr>")
        tag = "th" if ri == 0 else "td"
        for c in row:
            if not c.get("render"):
                continue
            a = ""
            if c["span"] > 1:
                a += f' colspan="{c["span"]}"'
            if c.get("rowspan", 1) > 1:
                a += f' rowspan="{c["rowspan"]}"'
            out.append(f"    <{tag}{a}>{cell_html(c['tc'], imgs)}</{tag}>")
        out.append("  </tr>")
    out.append("</table>")
    return "\n".join(out)


# ---------- paragraph rendering ----------

def render_paragraph(p) -> str:
    style = p.style.name or ""
    txt = clean(p.text).replace("\t", " — ")
    if not txt:
        return ""
    if style.startswith("Heading") or style == "Nadpis":
        m = re.search(r"(\d+)", style)
        level = int(m.group(1)) if m else 1
        return "#" * min(level + 1, 6) + " " + txt
    if style in CAPTION_STYLES or re.match(r"^(Figure|Table)\b", txt):
        return f"*{txt}*"
    if style in LIST_STYLES:
        return "- " + txt
    if style in NOTE_STYLES:
        return "> " + txt
    return txt


# ---------- main convert ----------

def convert(path, outdir):
    doc = docx.Document(path)
    docname = os.path.splitext(os.path.basename(path))[0]
    docdir = os.path.join(outdir, docname)
    imgs = ImgState(doc, docdir, docname)

    paras = [p for p in doc.paragraphs if p.text.strip()]
    annotation = clean(paras[1].text) if len(paras) > 1 else ""
    caption = next((clean(p.text) for p in paras
                    if clean(p.text).startswith("Published")), "")
    standard, name, parts = parse_name_table(doc.tables[0])
    note = next((clean(p.text) for p in paras
                 if clean(p.text).startswith("Note:")), "")

    fm = ["---"]
    m = CAPTION_RE.search(caption)
    if m:
        fm.append(f"published: {int(m.group('year'))}")
        if m.group("edition"):
            fm.append(f"edition: {int(m.group('edition'))}")
        fm.append(f"pages: {int(m.group('pages'))}")
    fm.append(f"title: {yaml_q(standard + ' - Extract')}")
    fm.append(f"standard: {yaml_q(standard)}")
    fm.append(f"name: {yaml_q(name)}")
    for i, part in enumerate(parts, 1):
        fm.append(f"name_{i}: {yaml_q(part)}")
    if annotation:
        fm.append(f"annotation: {yaml_q(annotation)}")
    if note:
        fm.append(f"note: {yaml_q(note)}")
    fm.append("---")

    # ordered body blocks
    blocks = []
    in_body = False
    for child in doc.element.body.iterchildren():
        if child.tag == P_TAG:
            p = Paragraph(child, doc)
            if not in_body:
                if p.text.strip().lower().startswith("introduction") \
                        and (p.style.name or "").startswith("Heading"):
                    in_body = True
                else:
                    continue
            blocks.append(("p", p))
        elif child.tag == TBL_TAG:
            if in_body:
                blocks.append(("tbl", Table(child, doc)))

    lines = []
    for i, (kind, obj) in enumerate(blocks):
        if kind == "tbl":
            lines.append(table_html(obj, imgs))
            continue
        rids = para_blip_rids(obj._p)
        for rid in rids:
            rel = imgs.save(rid)
            # alt text from neighbouring caption block
            alt = ""
            for j in (i + 1, i - 1):
                if 0 <= j < len(blocks) and blocks[j][0] == "p":
                    t = clean(blocks[j][1].text)
                    if re.match(r"^(Figure|Table)\b", t):
                        alt = t
                        break
            lines.append(f"![{alt}]({rel})")
        md = render_paragraph(obj)
        if md:
            lines.append(md)

    body = "\n\n".join(lines)
    os.makedirs(docdir, exist_ok=True)
    out_md = os.path.join(docdir, "index.md")
    with open(out_md, "w", encoding="utf-8") as fh:
        fh.write("\n".join(fm) + "\n\n" + body + "\n")
    return docname, imgs.n, sum(1 for k, _ in blocks if k == "tbl")


def main():
    indir, outdir = sys.argv[1], sys.argv[2]
    os.makedirs(outdir, exist_ok=True)
    for f in sorted(glob.glob(os.path.join(indir, "*.docx"))):
        if os.path.basename(f).startswith("~"):
            continue
        name, nimg, ntbl = convert(f, outdir)
        print(f"{name:24} figures={nimg} tables={ntbl}")


if __name__ == "__main__":
    main()
