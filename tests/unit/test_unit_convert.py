"""Unit tests for convert_docx.py functions."""
import re

import pytest
import docx

from convert_docx import clean, yaml_q, CAPTION_RE, render_paragraph


@pytest.mark.unit
def test_clean_whitespace():
    """Test clean() normalizes various whitespace characters."""
    # NBSP (non-breaking space)
    assert clean("hello\xa0world") == "hello world"
    
    # Multiple spaces
    assert clean("hello  world") == "hello world"
    assert clean("hello   world") == "hello world"
    
    # Tabs
    assert clean("hello\tworld") == "hello world"
    
    # Leading/trailing whitespace
    assert clean("  hello world  ") == "hello world"
    
    # Mixed whitespace
    assert clean("  hello\xa0\t  world  ") == "hello world"
    
    # Empty string
    assert clean("") == ""
    
    # Only whitespace
    assert clean("   ") == ""


@pytest.mark.unit
def test_yaml_q_escaping():
    """Test yaml_q() properly escapes YAML strings."""
    # Simple string
    assert yaml_q("hello") == '"hello"'
    
    # String with quotes
    assert yaml_q('say "hello"') == '"say \\"hello\\""'
    
    # String with backslashes
    assert yaml_q(r"path\to\file") == '"path\\\\to\\\\file"'
    
    # String with newlines (should be collapsed)
    assert yaml_q("line1\nline2") == '"line1 line2"'
    
    # String with multiple spaces (should be normalized)
    assert yaml_q("hello   world") == '"hello world"'
    
    # Mixed newlines and spaces
    assert yaml_q("hello\n  world\n  !") == '"hello world !"'
    
    # Empty string
    assert yaml_q("") == '""'


@pytest.mark.unit
def test_caption_re_parsing():
    """Test CAPTION_RE regex matches publication metadata."""
    # Full match with edition
    text1 = "Published in 2025 (Edition 2) by ISO. 123 pages total."
    m1 = CAPTION_RE.search(text1)
    assert m1 is not None
    assert m1.group("year") == "2025"
    assert m1.group("edition") == "2"
    assert m1.group("pages") == "123"
    
    # Without "in"
    text2 = "Published 2024 (Edition 3) with 456 pages."
    m2 = CAPTION_RE.search(text2)
    assert m2 is not None
    assert m2.group("year") == "2024"
    assert m2.group("edition") == "3"
    assert m2.group("pages") == "456"
    
    # Without edition
    text3 = "Published 2023 - total of 789 pages."
    m3 = CAPTION_RE.search(text3)
    assert m3 is not None
    assert m3.group("year") == "2023"
    assert m3.group("edition") is None
    assert m3.group("pages") == "789"
    
    # Case insensitive
    text4 = "PUBLISHED 2022 with 100 PAGES"
    m4 = CAPTION_RE.search(text4)
    assert m4 is not None
    assert m4.group("year") == "2022"
    
    # No match
    text5 = "This document has no publication info"
    m5 = CAPTION_RE.search(text5)
    assert m5 is None


@pytest.mark.unit
def test_caption_re_edge_cases():
    """Test CAPTION_RE edge cases."""
    # Edition with different capitalization
    text = "Published in 2025 (edition 1), 50 pages"
    m = CAPTION_RE.search(text)
    assert m is not None
    assert m.group("edition") == "1"
    
    # Multiple digit years and pages
    text2 = "Published 1999 (Edition 99) spanning 9999 pages"
    m2 = CAPTION_RE.search(text2)
    assert m2 is not None
    assert m2.group("year") == "1999"
    assert m2.group("edition") == "99"
    assert m2.group("pages") == "9999"


@pytest.mark.unit
def test_monospace_formatting():
    """Test that Courier New text is converted to inline code."""
    doc = docx.Document()
    para = doc.add_paragraph()
    
    # Add normal text
    run1 = para.add_run("registered under ")
    
    # Add monospace text
    run2 = para.add_run("{joint-iso-itut(2) its(28) gdd(5)}")
    run2.font.name = "Courier New"
    
    # Add more normal text
    run3 = para.add_run(" and the way")
    
    result = render_paragraph(para)
    assert "`{joint-iso-itut(2) its(28) gdd(5)}`" in result
    assert result == "registered under `{joint-iso-itut(2) its(28) gdd(5)}` and the way"


@pytest.mark.unit
def test_bold_formatting():
    """Test that bold text is converted to **bold**."""
    doc = docx.Document()
    para = doc.add_paragraph()
    run1 = para.add_run("The ")
    run2 = para.add_run("Graphic Data Dictionary (GDD)")
    run2.bold = True
    run3 = para.add_run(" is defined")
    
    result = render_paragraph(para)
    assert result == "The **Graphic Data Dictionary (GDD)** is defined"


@pytest.mark.unit
def test_italic_formatting():
    """Test that italic text is converted to *italic*."""
    doc = docx.Document()
    para = doc.add_paragraph()
    run1 = para.add_run("See ")
    run2 = para.add_run("ITS Terminology")
    run2.italic = True
    
    result = render_paragraph(para)
    assert result == "See *ITS Terminology*"


@pytest.mark.unit
def test_bold_italic_formatting():
    """Test that bold+italic text is converted to ***text***."""
    doc = docx.Document()
    para = doc.add_paragraph()
    run = para.add_run("This Extract does not replace...")
    run.bold = True
    run.italic = True
    
    result = render_paragraph(para)
    assert result == "***This Extract does not replace...***"


@pytest.mark.unit
def test_mixed_formatting_in_paragraph():
    """Test paragraph with multiple formatting types."""
    doc = docx.Document()
    para = doc.add_paragraph()
    
    r1 = para.add_run("Clause ")
    r2 = para.add_run("7.3")
    r2.bold = True
    r3 = para.add_run(" describes ")
    r4 = para.add_run("{joint-iso-itut(2)}")
    r4.font.name = "Courier New"
    r5 = para.add_run(" and ")
    r6 = para.add_run("Note")
    r6.italic = True
    
    result = render_paragraph(para)
    assert result == "Clause **7.3** describes `{joint-iso-itut(2)}` and *Note*"


@pytest.mark.unit
def test_code_ignores_other_formatting():
    """Test that code formatting takes precedence over bold/italic."""
    doc = docx.Document()
    para = doc.add_paragraph()
    
    # Add code text with bold+italic (should render as plain code)
    run = para.add_run("{1 2 7 xx}")
    run.font.name = "Courier New"
    run.bold = True
    run.italic = True
    
    result = render_paragraph(para)
    # Should be plain code, not ***`{1 2 7 xx}`***
    assert result == "`{1 2 7 xx}`"


@pytest.mark.unit
def test_whitespace_preservation():
    """Test that whitespace between runs is preserved."""
    doc = docx.Document()
    para = doc.add_paragraph()
    
    r1 = para.add_run("hello")
    r2 = para.add_run(" ")
    r3 = para.add_run("world")
    r3.bold = True
    
    result = render_paragraph(para)
    assert result == "hello **world**"
    
    # Test with space in formatted run
    para2 = doc.add_paragraph()
    p1 = para2.add_run("The ")
    p2 = para2.add_run("quick fox")
    p2.italic = True
    p3 = para2.add_run(" jumps")
    
    result2 = render_paragraph(para2)
    assert result2 == "The *quick fox* jumps"


@pytest.mark.unit
def test_spaces_outside_formatting():
    """Test that leading/trailing spaces are moved outside formatting markers."""
    doc = docx.Document()
    para = doc.add_paragraph()
    
    # Test trailing space in bold
    r1 = para.add_run("Clause ")
    r2 = para.add_run("8.2 ")
    r2.bold = True
    r3 = para.add_run("describes")
    
    result = render_paragraph(para)
    # Should be "Clause **8.2** describes" not "Clause **8.2 **describes"
    assert result == "Clause **8.2** describes"
    
    # Test leading space in italic
    para2 = doc.add_paragraph()
    p1 = para2.add_run("The")
    p2 = para2.add_run(" term")
    p2.italic = True
    p3 = para2.add_run(" here")
    
    result2 = render_paragraph(para2)
    assert result2 == "The *term* here"
    assert "* " not in result2 or result2.count("* ") == 1, "Space should be outside markers"
    
    # Test both leading and trailing spaces
    para3 = doc.add_paragraph()
    x1 = para3.add_run("word")
    x2 = para3.add_run(" middle ")
    x2.bold = True
    x3 = para3.add_run("word")
    
    result3 = render_paragraph(para3)
    assert result3 == "word **middle** word"
